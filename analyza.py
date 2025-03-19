import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import os
import logging
from docx import Document
from docx.shared import Inches
import re
from io import BytesIO

# Konfigurace loggeru
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Konstanty
OUTPUT_FOLDER = "output"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def sanitize_name(name: str) -> str:
    """Odstraní nepovolené znaky z textu, aby bylo možné bezpečně vytvářet názvy souborů."""
    return re.sub(r'[\\/*?:"<>|]', '_', name)

def load_data(file_path):
    logger.info(f"Načítám data ze souboru: {file_path}")
    excel_file = pd.ExcelFile(file_path)
    sheet_names = excel_file.sheet_names
    if "data" in sheet_names:
        df = pd.read_excel(file_path, sheet_name="data")
    else:
        df = pd.read_excel(file_path, sheet_name=sheet_names[0])
    for col in ["Jmeno", "Prijmeni", "Narozen"]:
        if col not in df.columns:
            raise KeyError(f"Chybí sloupec '{col}' v datech.")
    df["Identifikace"] = df["Jmeno"].astype(str) + " " + df["Prijmeni"].astype(str) + ", " + df["Narozen"].astype(str)
    return df

def format_val(val):
    if isinstance(val, pd.Timedelta):
        return f"{val.total_seconds():.2f}"
    try:
        return f"{float(val):.2f}"
    except (ValueError, TypeError):
        return str(val)

# Registrace fontů
base_dir = os.path.dirname(os.path.abspath(__file__))
times_font_path = os.path.join(base_dir, "times.ttf")
times_bold_font_path = os.path.join(base_dir, "timesbd.ttf")
pdfmetrics.registerFont(TTFont('TimesNewRoman', times_font_path))
pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', times_bold_font_path))

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="Custom-Regular", fontName="TimesNewRoman", fontSize=12))
styles.add(ParagraphStyle(name="Custom-Bold", fontName="TimesNewRoman-Bold", fontSize=14, spaceAfter=10, leading=16))

# Předdefinované skupiny grafů a popisky proměnných
GRAPH_GROUPS = [
    ("Poměr IR/ER", ["IR/ER (210°/s)", "IR/ER (300°/s)"], "ir_er_ratio.png"),
    ("Složení těla", ["Dominantni paze", "Dominantni paze - beztukova", "Dominantni noha", "Dominantni noha - beztukova", "Trupova hmotnost", "Trup - betukovy", "Beztukova hmota"], "body_composition.png"),
    ("Síla úchopu a rychlost podání", ["Sila uchopu", "Rychlost podani"], "grip_speed.png"),
    ("Vnitřní/Vnější rotace (210°/s)", ["Vnitrni rotace koncentricka (210°/s)", "Vnejsi rotace koncentricka (210°/s)", "Vnitrni rotace excentricka (210°/s)", "Vnejsi rotace excentricka (210°/s)"], "rotation_210.png"),
    ("Vnitřní/Vnější rotace (300°/s)", ["Vnitrni rotace koncentricka (300°/s)", "Vnejsi rotace koncentricka (300°/s)", "Vnitrni rotace excentricka (300°/s)", "Vnejsi rotace excentricka (300°/s)"], "rotation_300.png")
]

variable_legends = {
    "IR/ER (210°/s)": "Poměr vnitřní rotace k vnější rotaci při rychlosti 210°/s.",
    "IR/ER (300°/s)": "Poměr vnitřní rotace k vnější rotaci při rychlosti 300°/s.",
    "Dominantni paze": "Ukazuje sílu/velikost dominantní paže.",
    "Dominantni paze - beztukova": "Hodnota beztukové hmoty dominantní paže.",
    "Dominantni noha": "Ukazuje sílu/velikost dominantní nohy.",
    "Dominantni noha - beztukova": "Hodnota beztukové hmoty dominantní nohy.",
    "Trupova hmotnost": "Celková hmotnost trupu.",
    "Trup - betukovy": "Hmotnost trupu bez tukové složky.",
    "Beztukova hmota": "Celková beztuková hmota těla.",
    "Sila uchopu": "Síla úchopu, důležitá pro kontrolu rakety.",
    "Rychlost podani": "Rychlost podání, klíčová pro herní výkon.",
    "Vnitrni rotace koncentricka (210°/s)": "Izokinetická síla při vnitřní rotaci ramene, 210°/s.",
    "Vnejsi rotace koncentricka (210°/s)": "Izokinetická síla při vnější rotaci ramene, 210°/s.",
    "Vnitrni rotace excentricka (210°/s)": "Izokinetická síla při vnitřní rotaci ramene, 210°/s.",
    "Vnejsi rotace excentricka (210°/s)": "Izokinetická síla při vnější rotaci ramene, 210°/s.",
    "Vnitrni rotace koncentricka (300°/s)": "Izokinetická síla při vnitřní rotaci ramene, 300°/s.",
    "Vnejsi rotace koncentricka (300°/s)": "Izokinetická síla při vnější rotaci ramene, 300°/s.",
    "Vnitrni rotace excentricka (300°/s)": "Izokinetická síla při vnitřní rotaci ramene, 300°/s.",
    "Vnejsi rotace excentricka (300°/s)": "Izokinetická síla při vnější rotaci ramene, 300°/s."
}

def generate_graph(nazev, hodnoty_proband, hodnoty_avg, popisky, graph_type="bar", 
                   label_current="Aktuální měření", label_reference="Historické měření"):
    logger.info(f"Generuji graf: {nazev}, typ: {graph_type}")
    fig, ax = plt.subplots(figsize=(10, 6))
    if graph_type == "bar":
        x = np.arange(len(popisky))
        bar_width = 0.4
        bars1 = ax.bar(x - bar_width/2, hodnoty_proband, bar_width, label=label_current, color="#1F4E79", alpha=0.9, edgecolor="black")
        bars2 = ax.bar(x + bar_width/2, hodnoty_avg, bar_width, label=label_reference, color="#A0A0A0", alpha=0.9, edgecolor="black")
        for bar in bars1 + bars2:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2, height + 0.1, f"{height:.2f}", ha="center", va="bottom", fontsize=12, fontweight="bold")
        ax.set_title(nazev, fontsize=16, fontweight="bold", pad=20)
        ax.set_xticks(x)
        ax.set_xticklabels(popisky, rotation=20, ha="right", fontsize=12)
        ax.legend(fontsize=12)
        ax.yaxis.grid(True, linestyle="--", alpha=0.7)
    elif graph_type == "line":
        ax.plot(popisky, hodnoty_proband, marker="o", label=label_current, color="#1F4E79")
        ax.plot(popisky, hodnoty_avg, marker="o", label=label_reference, color="#A0A0A0")
        for i, y in enumerate(hodnoty_proband):
            ax.text(i, y + 0.1, f"{y:.2f}", ha="center", va="bottom", fontsize=12, fontweight="bold")
        ax.set_title(nazev, fontsize=16, fontweight="bold", pad=20)
        ax.set_xticks(range(len(popisky)))
        ax.set_xticklabels(popisky, rotation=20, ha="right", fontsize=12)
        ax.legend(fontsize=12)
        ax.yaxis.grid(True, linestyle="--", alpha=0.7)
    elif graph_type == "scatter":
        ax.scatter(range(len(popisky)), hodnoty_proband, label=label_current, color="#1F4E79")
        ax.scatter(range(len(popisky)), hodnoty_avg, label=label_reference, color="#A0A0A0")
        for i, y in enumerate(hodnoty_proband):
            ax.text(i, y + 0.1, f"{y:.2f}", ha="center", va="bottom", fontsize=12, fontweight="bold")
        ax.set_title(nazev, fontsize=16, fontweight="bold", pad=20)
        ax.set_xticks(range(len(popisky)))
        ax.set_xticklabels(popisky, rotation=20, ha="right", fontsize=12)
        ax.legend(fontsize=12)
        ax.yaxis.grid(True, linestyle="--", alpha=0.7)
    else:
        logger.warning(f"Neznámý typ grafu: {graph_type}, používám 'bar'.")
        return generate_graph(nazev, hodnoty_proband, hodnoty_avg, popisky, graph_type="bar",
                              label_current=label_current, label_reference=label_reference)
    
    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches="tight")
    plt.close()
    buf.seek(0)
    return buf

def interpretuj_graf(nazev, hodnoty_proband, hodnoty_avg, popisky):
    desired_direction = {
        "Vnitrni rotace koncentricka (210°/s)": "higher",
        "Vnejsi rotace koncentricka (210°/s)": "higher",
        "Vnitrni rotace excentricka (210°/s)": "higher",
        "Vnejsi rotace excentricka (210°/s)": "higher",
        "Vnitrni rotace koncentricka (300°/s)": "higher",
        "Vnejsi rotace koncentricka (300°/s)": "higher",
        "Vnitrni rotace excentricka (300°/s)": "higher",
        "Vnejsi rotace excentricka (300°/s)": "higher",
        "Rychlost podani": "higher",
        "Sila uchopu": "higher",
        "Dominantni paze": "higher",
        "Dominantni noha": "higher",
        "Trupova hmotnost": "optimal",
        "Telesny tuk": "lower",
        "Dominantni paze - beztukova": "higher",
        "Dominantni noha - beztukova": "higher",
        "Trup - betukovy": "optimal",
        "Beztukova hmota": "higher",
        "IR/ER (210°/s)": "optimal",
        "IR/ER (300°/s)": "optimal"
    }
    
    interpretations = []
    for i, label in enumerate(popisky):
        diff = hodnoty_proband[i] - hodnoty_avg[i]
        direction = desired_direction.get(label, "higher")
        if abs(diff) < 0.1:
            interpretations.append(f"U '{label}' je hodnota aktuálního měření srovnatelná s referenční hodnotou.")
        else:
            if direction == "higher":
                if diff > 0:
                    interpretations.append(f"U '{label}' je aktuální měření o {diff:.2f} vyšší, což značí zlepšení.")
                else:
                    interpretations.append(f"U '{label}' je aktuální měření o {abs(diff):.2f} nižší, což může naznačovat potřebu zlepšení.")
            elif direction == "lower":
                if diff < 0:
                    interpretations.append(f"U '{label}' je aktuální měření o {abs(diff):.2f} nižší, což značí zlepšení.")
                else:
                    interpretations.append(f"U '{label}' je aktuální měření o {diff:.2f} vyšší, což může být nežádoucí.")
            else:
                interpretations.append(f"U '{label}' je aktuální měření o {abs(diff):.2f} odlišné od referenční hodnoty.")
    return " ".join(interpretations)

def generuj_analyzu(proband_id, file_path, zaverecne_hodnoceni=None,
                     selected_columns=None, selected_graphs=None,
                     selected_graph_type="bar", data_df=None, comparison_data=None,
                     advanced_stats=False, group_label=None, selected_graph_vars=None):
    logger.info(f"Generuji analýzu pro probanda: {proband_id}")
    if data_df is not None:
        df = data_df.copy()
    else:
        df = load_data(file_path)
    df.columns = df.columns.str.strip().str.replace("\\s+", " ", regex=True)
    
    default_columns = ["Jmeno", "Prijmeni", "Narozen", "Identifikace", "Vek", "Vyska", "Hmotnost"]
    if selected_columns is None:
        selected_columns = [col for col in df.columns if col not in default_columns]
    
    selected_columns = [c for c in selected_columns if c in df.columns]
    
    if "Vnitrni rotace koncentricka (210°/s)" in df.columns and "Vnejsi rotace koncentricka (210°/s)" in df.columns:
        df["IR/ER (210°/s)"] = df["Vnitrni rotace koncentricka (210°/s)"] / df["Vnejsi rotace koncentricka (210°/s)"]
        if "IR/ER (210°/s)" not in selected_columns:
            selected_columns.append("IR/ER (210°/s)")
    if "Vnitrni rotace koncentricka (300°/s)" in df.columns and "Vnejsi rotace koncentricka (300°/s)" in df.columns:
        df["IR/ER (300°/s)"] = df["Vnitrni rotace koncentricka (300°/s)"] / df["Vnejsi rotace koncentricka (300°/s)"]
        if "IR/ER (300°/s)" not in selected_columns:
            selected_columns.append("IR/ER (300°/s)")
    
    df.fillna(0, inplace=True)
    proband_data = df[df["Identifikace"] == proband_id].iloc[0]
    
    pdf_path = os.path.join(OUTPUT_FOLDER, f"analyza_{sanitize_name(proband_id)}.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    elements = []
    
    elements.append(Paragraph("Univerzita Karlova, Fakulta tělesné výchovy a sportu", styles["Custom-Bold"]))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph(f"Analýza probanda {proband_id}", styles["Custom-Bold"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"Věk: {proband_data['Vek']} let", styles["Custom-Regular"]))
    elements.append(Paragraph(f"Výška: {proband_data['Vyska']} cm", styles["Custom-Regular"]))
    elements.append(Paragraph(f"Hmotnost: {proband_data['Hmotnost']} kg", styles["Custom-Regular"]))
    elements.append(Spacer(1, 12))
    
    if comparison_data is None:
        date_value = proband_data.get("DatumMereni", "N/A")
        if group_label is not None:
            if group_label == "Aktuální skupina":
                elements.append(Paragraph(f"Porovnání probanda s průměrnou hodnotou aktuální skupiny – datum: {date_value}", styles["Custom-Regular"]))
            else:
                elements.append(Paragraph(f"Porovnání probanda s průměrnou hodnotou vybrané populace – datum: {date_value}", styles["Custom-Regular"]))
        else:
            elements.append(Paragraph(f"Porovnání probanda s průměrným výsledkem skupiny – datum: {date_value}", styles["Custom-Regular"]))
    else:
        current_date = proband_data.get("DatumMereni", "N/A")
        historical_date = comparison_data.get("DatumMereni", "N/A")
        elements.append(Paragraph(f"Datum aktuálního měření: {current_date}    Datum vybraného historického měření: {historical_date}", styles["Custom-Regular"]))
    
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph("Výsledky měření", styles["Custom-Bold"]))
    if comparison_data is None:
        data_table = []
        header = ["Parametr", "Aktuální", "Průměr", "Rozdíl"]
        data_table.append(header)
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                prumer = df[col].mean()
                rozdil = proband_data[col] - prumer
                data_table.append([col, format_val(proband_data[col]), format_val(prumer), format_val(rozdil)])
        table = Table(data_table, hAlign='LEFT')
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'TimesNewRoman-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table)
    else:
        data_table = []
        header = ["Parametr", "Aktuální", "Historické", "Rozdíl"]
        data_table.append(header)
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                current_val = proband_data[col]
                hist_val = comparison_data.get(col, None)
                if hist_val is None and col == "IR/ER (210°/s)":
                    if ("Vnitrni rotace koncentricka (210°/s)" in comparison_data and "Vnejsi rotace koncentricka (210°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (210°/s)"] / comparison_data["Vnejsi rotace koncentricka (210°/s)"]
                if hist_val is None and col == "IR/ER (300°/s)":
                    if ("Vnitrni rotace koncentricka (300°/s)" in comparison_data and "Vnejsi rotace koncentricka (300°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (300°/s)"] / comparison_data["Vnejsi rotace koncentricka (300°/s)"]
                if hist_val is not None:
                    diff_val = current_val - hist_val
                    data_table.append([col, format_val(current_val), format_val(hist_val), format_val(diff_val)])
        table = Table(data_table, hAlign='LEFT')
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'TimesNewRoman-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table)
    
    elements.append(Spacer(1, 12))
    
    if comparison_data is not None:
        elements.append(Paragraph("• Porovnání: Aktuální měření vs. historické měření.", styles["Custom-Regular"]))
    else:
        if group_label is not None:
            elements.append(Paragraph(f"• Porovnání: Proband vs. {group_label}.", styles["Custom-Regular"]))
        else:
            elements.append(Paragraph("• Porovnání: Proband vs. průměr skupiny.", styles["Custom-Regular"]))
    elements.append(Spacer(1, 12))
    
    if advanced_stats:
        elements.append(Paragraph("Rozšířené statistiky (vypočteno z aktuálních měření)", styles["Custom-Bold"]))
        numeric_cols = [c for c in selected_columns if c in df.columns and pd.api.types.is_numeric_dtype(df[c])]
        table_data = [["Parametr", "Medián", "Nejlepší", "Nejhorší", "CI (spodní)", "CI (horní)"]]
        for col in numeric_cols:
            data = df[col].values
            median_val = np.median(data)
            best_val = np.max(data)
            worst_val = np.min(data)
            ci_lower = np.percentile(data, 2.5)
            ci_upper = np.percentile(data, 97.5)
            table_data.append([col, format_val(median_val), format_val(best_val), format_val(worst_val), format_val(ci_lower), format_val(ci_upper)])
        table2 = Table(table_data, hAlign='LEFT')
        table2.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table2)
        elements.append(Spacer(1, 12))
    
    from analyza import generate_graph, interpretuj_graf
    for nazev, popisky, _ in GRAPH_GROUPS:
        if selected_graphs is not None and nazev not in selected_graphs:
            continue
        filtered_popisky = [p for p in popisky if p in selected_columns]
        if not filtered_popisky:
            continue
        if comparison_data is None:
            comp_values = [df[p].mean() for p in filtered_popisky]
            if group_label is not None:
                current_label = "Proband"
                reference_label = group_label
            else:
                current_label = "Proband"
                reference_label = "Průměr skupiny"
        else:
            comp_values = []
            current_label = "Aktuální měření"
            reference_label = "Historické měření"
            for p in filtered_popisky:
                val = comparison_data.get(p, None)
                if val is None and p == "IR/ER (210°/s)":
                    if ("Vnitrni rotace koncentricka (210°/s)" in comparison_data and "Vnejsi rotace koncentricka (210°/s)" in comparison_data):
                        val = comparison_data["Vnitrni rotace koncentricka (210°/s)"] / comparison_data["Vnejsi rotace koncentricka (210°/s)"]
                if val is None and p == "IR/ER (300°/s)":
                    if ("Vnitrni rotace koncentricka (300°/s)" in comparison_data and "Vnejsi rotace koncentricka (300°/s)" in comparison_data):
                        val = comparison_data["Vnitrni rotace koncentricka (300°/s)"] / comparison_data["Vnejsi rotace koncentricka (300°/s)"]
                comp_values.append(val)
        graph_img = generate_graph(nazev,
                    [proband_data[p] for p in filtered_popisky],
                    comp_values,
                    filtered_popisky,
                    graph_type=selected_graph_type,
                    label_current=current_label,
                    label_reference=reference_label)
        elements.append(PageBreak())
        elements.append(Paragraph(nazev, styles["Custom-Bold"]))
        elements.append(Image(graph_img, width=450, height=300))
        legend_text = "Legenda:\n"
        for var in filtered_popisky:
            if var in variable_legends:
                legend_text += f"{var}: {variable_legends[var]}\n"
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(legend_text, styles["Custom-Regular"]))
        interpretation_text = interpretuj_graf(nazev,
                                                 [proband_data[p] for p in filtered_popisky],
                                                 comp_values,
                                                 filtered_popisky)
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Vyhodnocení grafu:", styles["Custom-Bold"]))
        elements.append(Paragraph(interpretation_text, styles["Custom-Regular"]))
    
    if selected_graph_vars is not None:
        for var in selected_graph_vars:
            if comparison_data is None:
                avg_val = df[var].mean()
                label_ref = "Průměr skupiny" if group_label is None else group_label
                current_label = "Proband"
            else:
                avg_val = comparison_data.get(var, 0)
                label_ref = "Historické měření"
                current_label = "Aktuální měření"
            graph_img = generate_graph(var,
                        [proband_data[var]],
                        [avg_val],
                        [var],
                        graph_type=selected_graph_type,
                        label_current=current_label,
                        label_reference=label_ref)
            elements.append(PageBreak())
            elements.append(Paragraph(var, styles["Custom-Bold"]))
            elements.append(Image(graph_img, width=450, height=300))
            legend_text = f"Legenda: Graf proměnné {var} zobrazuje hodnotu probanda (viz {current_label}) a průměr skupiny/historické měření (viz {label_ref})."
            elements.append(Spacer(1, 12))
            elements.append(Paragraph(legend_text, styles["Custom-Regular"]))
            evaluation = interpretuj_graf(var, [proband_data[var]], [avg_val], [var])
            elements.append(Spacer(1, 12))
            elements.append(Paragraph("Vyhodnocení grafu:", styles["Custom-Bold"]))
            elements.append(Paragraph(evaluation, styles["Custom-Regular"]))
    
    if zaverecne_hodnoceni and zaverecne_hodnoceni.strip():
        elements.append(PageBreak())
        elements.append(Paragraph("Závěrečné doporučení", styles["Custom-Bold"]))
        elements.append(Spacer(1, 12))
        for para in zaverecne_hodnoceni.strip().split("\n\n"):
            p = Paragraph(para.strip().replace("\n", "<br/>"), styles["Custom-Regular"])
            elements.append(p)
            elements.append(Spacer(1, 12))
    
    doc.build(elements)
    logger.info(f"PDF report vygenerován: {pdf_path}")
    return pdf_path

def generuj_word_report(proband_id, file_path, zaverecne_hodnoceni=None,
                        selected_columns=None, selected_graphs=None,
                        selected_graph_type="bar",  # parametr přidaný
                        advanced_stats=False, group_label=None, data_df=None, comparison_data=None,
                        selected_graph_vars=None):
    if data_df is not None:
        df = data_df.copy()
    else:
        df = load_data(file_path)
    df.columns = df.columns.str.strip().str.replace("\\s+", " ", regex=True)
    default_columns = ["Jmeno", "Prijmeni", "Narozen", "Identifikace", "Vek", "Vyska", "Hmotnost"]
    if selected_columns is None:
        selected_columns = [col for col in df.columns if col not in default_columns]
    df.fillna(0, inplace=True)
    proband_data = df[df["Identifikace"] == proband_id].iloc[0]
    
    document = Document()
    document.add_heading("Univerzita Karlova, Fakulta tělesné výchovy a sportu", level=1)
    document.add_heading(f"Analýza probanda {proband_id}", level=2)
    document.add_paragraph(f"Věk: {proband_data['Vek']} let")
    document.add_paragraph(f"Výška: {proband_data['Vyska']} cm")
    document.add_paragraph(f"Hmotnost: {proband_data['Hmotnost']} kg")
    
    if comparison_data is None:
        date_value = proband_data.get("DatumMereni", "N/A")
        if group_label is not None:
            if group_label == "Aktuální skupina":
                document.add_paragraph(f"Porovnání probanda s průměrnou hodnotou aktuální skupiny – datum: {date_value}")
            else:
                document.add_paragraph(f"Porovnání probanda s průměrnou hodnotou vybrané populace – datum: {date_value}")
        else:
            document.add_paragraph(f"Porovnání probanda s průměrným výsledkem skupiny – datum: {date_value}")
    else:
        current_date = proband_data.get("DatumMereni", "N/A")
        historical_date = comparison_data.get("DatumMereni", "N/A")
        document.add_paragraph(f"Datum aktuálního měření: {current_date}    Datum vybraného historického měření: {historical_date}")
    
    document.add_heading("Výsledky měření", level=3)
    if comparison_data is None:
        table_data = [["Parametr", "Aktuální", "Průměr", "Rozdíl"]]
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                prumer = df[col].mean()
                rozdil = proband_data[col] - prumer
                table_data.append([col, format_val(proband_data[col]), format_val(prumer), format_val(rozdil)])
    else:
        table_data = [["Parametr", "Aktuální", "Historické", "Rozdíl"]]
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                current_val = proband_data[col]
                hist_val = comparison_data.get(col, None)
                if hist_val is None and col == "IR/ER (210°/s)":
                    if ("Vnitrni rotace koncentricka (210°/s)" in comparison_data and "Vnejsi rotace koncentricka (210°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (210°/s)"] / comparison_data["Vnejsi rotace koncentricka (210°/s)"]
                if hist_val is None and col == "IR/ER (300°/s)":
                    if ("Vnitrni rotace koncentricka (300°/s)" in comparison_data and "Vnejsi rotace koncentricka (300°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (300°/s)"] / comparison_data["Vnejsi rotace koncentricka (300°/s)"]
                if hist_val is not None:
                    diff_val = current_val - hist_val
                    table_data.append([col, format_val(current_val), format_val(hist_val), format_val(diff_val)])
    table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)
    
    if advanced_stats:
        document.add_heading("Rozšířené statistiky (aktuální měření)", level=3)
        adv_table_data = [["Parametr", "Medián", "Nejlepší", "Nejhorší", "CI (spodní)", "CI (horní)"]]
        numeric_cols = [c for c in selected_columns if c in df.columns and pd.api.types.is_numeric_dtype(df[c])]
        for col in numeric_cols:
            data = df[col].values
            median_val = np.median(data)
            best_val = np.max(data)
            worst_val = np.min(data)
            ci_lower = np.percentile(data, 2.5)
            ci_upper = np.percentile(data, 97.5)
            adv_table_data.append([col, format_val(median_val), format_val(best_val), format_val(worst_val), format_val(ci_lower), format_val(ci_upper)])
        adv_table = document.add_table(rows=len(adv_table_data), cols=len(adv_table_data[0]))
        for i, row in enumerate(adv_table_data):
            for j, cell in enumerate(row):
                adv_table.cell(i, j).text = str(cell)
    
    if selected_graphs is not None:
        for nazev, popisky, _ in GRAPH_GROUPS:
            if nazev not in selected_graphs:
                continue
            filtered_popisky = [p for p in popisky if p in selected_columns]
            if not filtered_popisky:
                continue
            if comparison_data is None:
                comp_values = [df[p].mean() for p in filtered_popisky]
                current_label = "Proband"
                reference_label = "Průměr skupiny"
            else:
                comp_values = []
                current_label = "Aktuální měření"
                reference_label = "Historické měření"
                for p in filtered_popisky:
                    val = comparison_data.get(p, None)
                    if val is None and p == "IR/ER (210°/s)":
                        if ("Vnitrni rotace koncentricka (210°/s)" in comparison_data and "Vnejsi rotace koncentricka (210°/s)" in comparison_data):
                            val = comparison_data["Vnitrni rotace koncentricka (210°/s)"] / comparison_data["Vnejsi rotace koncentricka (210°/s)"]
                    if val is None and p == "IR/ER (300°/s)":
                        if ("Vnitrni rotace koncentricka (300°/s)" in comparison_data and "Vnejsi rotace koncentricka (300°/s)" in comparison_data):
                            val = comparison_data["Vnitrni rotace koncentricka (300°/s)"] / comparison_data["Vnejsi rotace koncentricka (300°/s)"]
                    comp_values.append(val)
            graph_img = generate_graph(nazev,
                        [proband_data[p] for p in filtered_popisky],
                        comp_values,
                        filtered_popisky,
                        graph_type=selected_graph_type,
                        label_current=current_label,
                        label_reference=reference_label)
            document.add_heading(nazev, level=3)
            document.add_picture(graph_img, width=Inches(6))
            legend_text = "Legenda:\n"
            for var in filtered_popisky:
                if var in variable_legends:
                    legend_text += f"{var}: {variable_legends[var]}\n"
            document.add_paragraph(legend_text)
            document.add_paragraph("Vyhodnocení grafu:")
            interpretation_text = interpretuj_graf(nazev,
                                                     [proband_data[p] for p in filtered_popisky],
                                                     [df[p].mean() for p in filtered_popisky] if comparison_data is None else [comparison_data.get(p, None) for p in filtered_popisky],
                                                     filtered_popisky)
            document.add_paragraph(interpretation_text)
    
    if selected_graph_vars is not None:
        for var in selected_graph_vars:
            if comparison_data is None:
                avg_val = df[var].mean()
                label_ref = "Průměr skupiny" if group_label is None else group_label
                current_label = "Proband"
            else:
                avg_val = comparison_data.get(var, 0)
                label_ref = "Historické měření"
                current_label = "Aktuální měření"
            graph_img = generate_graph(var,
                        [proband_data[var]],
                        [avg_val],
                        [var],
                        graph_type=selected_graph_type,
                        label_current=current_label,
                        label_reference=label_ref)
            document.add_heading(var, level=3)
            document.add_picture(graph_img, width=Inches(6))
            legend_text = f"Legenda: Graf proměnné {var} zobrazuje hodnotu probanda (viz {current_label}) a průměr skupiny/historické měření (viz {label_ref})."
            document.add_paragraph(legend_text)
            evaluation = interpretuj_graf(var, [proband_data[var]], [avg_val], [var])
            document.add_paragraph("Vyhodnocení grafu:")
            document.add_paragraph(evaluation)
    
    if zaverecne_hodnoceni and zaverecne_hodnoceni.strip():
        document.add_heading("Závěrečné doporučení", level=3)
        for para in zaverecne_hodnoceni.strip().split("\n\n"):
            document.add_paragraph(para.strip())
    
    word_path = os.path.join(OUTPUT_FOLDER, f"analyza_{sanitize_name(proband_id)}.docx")
    document.save(word_path)
    return word_path

def priprav_podklad(proband_id, file_path, selected_columns=None, data_df=None, comparison_data=None):
    logger.info("Připravuji textový podklad pro GPT.")
    if data_df is not None:
        df = data_df.copy()
    else:
        df = load_data(file_path)
    df.columns = df.columns.str.strip().str.replace("\\s+", " ", regex=True)
    
    default_columns = ["Jmeno", "Prijmeni", "Narozen", "Identifikace", "Vek", "Vyska", "Hmotnost"]
    if selected_columns is None:
        selected_columns = [col for col in df.columns if col not in default_columns]
    
    if "Vnitrni rotace koncentricka (210°/s)" in df.columns and "Vnejsi rotace koncentricka (210°/s)" in df.columns:
        df["IR/ER (210°/s)"] = df["Vnitrni rotace koncentricka (210°/s)"] / df["Vnejsi rotace koncentricka (210°/s)"]
        if "IR/ER (210°/s)" not in selected_columns:
            selected_columns.append("IR/ER (210°/s)")
    if "Vnitrni rotace koncentricka (300°/s)" in df.columns and "Vnejsi rotace koncentricka (300°/s)" in df.columns:
        df["IR/ER (300°/s)"] = df["Vnitrni rotace koncentricka (300°/s)"] / df["Vnejsi rotace koncentricka (300°/s)"]
        if "IR/ER (300°/s)" not in selected_columns:
            selected_columns.append("IR/ER (300°/s)")
    
    df.fillna(0, inplace=True)
    proband_data = df[df["Identifikace"] == proband_id].iloc[0]
    
    podklad = []
    podklad.append(f"Podklad pro hodnocení probanda {proband_id}")
    podklad.append("-" * 50)
    podklad.append(f"Věk: {proband_data['Vek']} let")
    podklad.append(f"Výška: {proband_data['Vyska']} cm")
    podklad.append(f"Hmotnost: {proband_data['Hmotnost']} kg")
    podklad.append("")
    podklad.append("Výsledky měření:")
    
    if comparison_data is None:
        header = f"{'Parametr':30} {'Aktuální':>10} {'Průměr':>10} {'Rozdíl':>10}"
        podklad.append(header)
        podklad.append("-" * len(header))
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                prumer = df[col].mean()
                rozdil = proband_data[col] - prumer
                podklad.append(f"{col:30} {format_val(proband_data[col]):>10} {format_val(prumer):>10} {format_val(rozdil):>10}")
    else:
        header = f"{'Parametr':30} {'Aktuální':>10} {'Historické':>10} {'Rozdíl':>10}"
        podklad.append(header)
        podklad.append("-" * len(header))
        for col in df.columns:
            if col not in default_columns and col in selected_columns and pd.api.types.is_numeric_dtype(df[col]):
                current_val = proband_data[col]
                hist_val = comparison_data.get(col, None)
                if hist_val is None and col == "IR/ER (210°/s)":
                    if ("Vnitrni rotace koncentricka (210°/s)" in comparison_data and "Vnejsi rotace koncentricka (210°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (210°/s)"] / comparison_data["Vnejsi rotace koncentricka (210°/s)"]
                if hist_val is None and col == "IR/ER (300°/s)":
                    if ("Vnitrni rotace koncentricka (300°/s)" in comparison_data and "Vnejsi rotace koncentricka (300°/s)" in comparison_data):
                        hist_val = comparison_data["Vnitrni rotace koncentricka (300°/s)"] / comparison_data["Vnejsi rotace koncentricka (300°/s)"]
                if hist_val is not None:
                    diff_val = current_val - hist_val
                    podklad.append(f"{col:30} {format_val(current_val):>10} {format_val(hist_val):>10} {format_val(diff_val):>10}")
    
    podklad.append("")
    podklad.append("Instrukce:")
    podklad.append("Na základě těchto výsledků vygenerujte prosím závěrečné hodnocení, které obsahuje:")
    podklad.append("- Celkové zhodnocení probanda")
    podklad.append("- Hlavní body (silné a slabé stránky)")
    podklad.append("- Doporučení pro trénink: koncentrický, excentrický, izometrický, plyometrický")
    podklad.append("- Návrh konkrétních cviků")
    podklad.append("- Komentář založený na vědeckých článcích o tenise")
    
    return "\n".join(podklad)
